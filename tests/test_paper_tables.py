"""The publication tables must be readable evidence, not a nicer-looking claim.

The properties asserted here are the ones a reviewer would check by hand: that
every number traces to a committed file, that a missing measurement says so
instead of reading as zero, that failed cases are still listed, and that the
illustrative example does not appear in the collection table.
"""
from __future__ import annotations

import csv
import json
import subprocess
import sys
from pathlib import Path

import pytest


# python-docx is in the `paper` extra rather than `test`, so a
# checkout that installs only the test extra skips these instead of
# reporting a missing module as a failure of the manuscript.
pytest.importorskip(
    "docx",
    reason="python-docx is in the `paper` extra; install -e \".[paper]\" to check the manuscript")

REPO_ROOT = Path(__file__).resolve().parents[1]
TABLES = REPO_ROOT / "paper_results" / "tables"
BUILDER = REPO_ROOT / "tools" / "tables" / "build_paper_tables.py"

sys.path.insert(0, str(REPO_ROOT / "tools" / "tables"))
from build_paper_tables import (  # noqa: E402
    BUILDERS, ILLUSTRATIVE, UNAVAILABLE,
)


def _sources() -> list[Path]:
    return sorted(TABLES.glob("table*.csv"))


def test_every_table_was_generated():
    assert len(_sources()) == len(BUILDERS)


@pytest.mark.parametrize("path", _sources(), ids=lambda p: p.stem)
def test_each_table_has_a_provenance_sidecar_naming_hashed_inputs(path: Path):
    sidecar = path.with_name(f"{path.stem}_provenance.json")
    assert sidecar.is_file(), f"no provenance beside {path.name}"
    record = json.loads(sidecar.read_text(encoding="utf-8"))
    assert record["command"].startswith("python tools/tables/")
    assert record["row_count"] > 0
    for entry in record["inputs"]:
        assert (REPO_ROOT / entry["path"]).is_file(), entry["path"]
        assert len(entry["sha256"]) == 64


@pytest.mark.parametrize("path", _sources(), ids=lambda p: p.stem)
def test_each_table_has_a_json_twin_with_the_same_rows(path: Path):
    twin = path.with_suffix(".json")
    assert twin.is_file()
    record = json.loads(twin.read_text(encoding="utf-8"))
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.reader(handle))
    assert rows[0] == record["columns"]
    assert len(rows) - 1 == len(record["rows"])


def test_the_word_document_holds_real_tables_not_pictures():
    from docx import Document

    document = Document(str(TABLES / "paper_tables.docx"))
    assert len(document.tables) == len(BUILDERS)
    for table in document.tables:
        assert len(table.rows) > 1
        assert any(cell.text.strip() for cell in table.rows[1].cells)
    assert not document.inline_shapes, "a table was embedded as an image"


def test_the_collection_table_excludes_the_illustrative_example():
    path = TABLES / "table6_collection_parameter_sensitivity.csv"
    with path.open(newline="", encoding="utf-8") as handle:
        models = {row["Model"] for row in csv.DictReader(handle)}
    assert not (models & ILLUSTRATIVE), (
        "the illustrative example must not appear in the collection table")
    assert len(models) >= 15


def test_failed_and_unresolved_cases_stay_in_the_tables():
    """A denominator that quietly loses its failures is not a denominator."""
    with (TABLES / "table2_abaqus_paired_validation.csv").open(
            newline="", encoding="utf-8") as handle:
        overall = [row["Overall"] for row in csv.DictReader(handle)]
    assert any(status != "passed" for status in overall)

    with (TABLES / "table6_collection_parameter_sensitivity.csv").open(
            newline="", encoding="utf-8") as handle:
        verdicts = [row["Verdict"] for row in csv.DictReader(handle)]
    assert "unresolved" in verdicts


def test_absent_measurements_say_so_rather_than_reading_as_zero():
    found = False
    for path in _sources():
        text = path.read_text(encoding="utf-8")
        if UNAVAILABLE in text:
            found = True
    assert found, (
        "no table marks an unavailable quantity; either every quantity was "
        "measured, or a missing one is being rendered as something else")


def test_no_table_names_a_home_directory():
    for path in _sources() + list(TABLES.glob("table*.json")):
        assert "/home/" not in path.read_text(encoding="utf-8"), path.name


def test_regenerates_identically_except_for_its_timestamps(tmp_path):
    """A table that changes without the evidence changing is not evidence."""
    done = subprocess.run(
        [sys.executable, str(BUILDER), "--out-dir", str(tmp_path)],
        capture_output=True, text=True, cwd=str(REPO_ROOT))
    assert done.returncode == 0, done.stdout + done.stderr
    for path in _sources():
        rebuilt = tmp_path / path.name
        assert rebuilt.is_file(), path.name
        assert rebuilt.read_text(encoding="utf-8") == \
            path.read_text(encoding="utf-8"), f"{path.name} changed"
