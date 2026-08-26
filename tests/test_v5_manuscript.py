"""The manuscript must not be able to state a number the evidence does not.

The substitution map raises on a missing key, so a claim without a measurement
behind it fails the build rather than reaching a page. These tests pin that
down, together with the limits SoftwareX imposes and the separation from V4.
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = REPO_ROOT / "docs" / "manuscript" / "UMAT_OTI_SoftwareX_V5.docx"
PROVENANCE = MANUSCRIPT.with_name(f"{MANUSCRIPT.stem}_provenance.json")
BUILDER = REPO_ROOT / "tools" / "manuscript" / "build_v5_manuscript.py"

sys.path.insert(0, str(REPO_ROOT / "tools" / "manuscript"))
from build_v5_manuscript import (  # noqa: E402
    SECTIONS, WORD_LIMIT, StrictValues, _derived,
)
from evidence_values import collect  # noqa: E402


def _record() -> dict:
    return json.loads(PROVENANCE.read_text(encoding="utf-8"))


def test_the_manuscript_and_its_provenance_exist():
    assert MANUSCRIPT.is_file()
    assert PROVENANCE.is_file()


def test_it_is_within_the_softwarex_limits():
    record = _record()
    assert record["word_count"] <= WORD_LIMIT
    assert record["figure_count"] <= 6


def test_every_substituted_value_names_a_committed_file():
    for key, entry in _record()["substituted_values"].items():
        assert (REPO_ROOT / entry["source"]).is_file(), f"{key}: {entry['source']}"


def test_a_claim_without_a_measurement_fails_the_build():
    """This is the property the whole approach rests on."""
    with pytest.raises(KeyError, match="does not provide"):
        "{a_number_nobody_measured}".format_map(StrictValues({}))


def test_every_placeholder_in_the_text_is_provided_by_the_evidence():
    available = set(_derived(collect()))
    for heading, paragraphs in SECTIONS:
        for paragraph in paragraphs:
            for key in re.findall(r"\{(\w+)\}", paragraph):
                assert key in available, f"{heading} cites unmeasured {key!r}"


def test_no_placeholder_survived_into_the_document():
    from docx import Document

    text = "\n".join(p.text for p in Document(str(MANUSCRIPT)).paragraphs)
    assert not re.search(r"\{\w+\}", text), "an unsubstituted placeholder"
    assert "TODO" not in text and "TBD" not in text


def test_the_figures_are_the_real_ones():
    from docx import Document

    document = Document(str(MANUSCRIPT))
    assert len(document.inline_shapes) == _record()["figure_count"]
    for name in _record()["figures"]:
        assert (REPO_ROOT / "paper_results" / "figures" / name).is_file(), name


def test_it_does_not_read_or_write_v4():
    source = BUILDER.read_text(encoding="utf-8")
    assert "V4" not in source.replace("V4 is not read", "")


def test_withdrawn_claims_do_not_reappear():
    """Figures the present release does not reproduce must stay out.

    These were in the earlier revision and are not reproduced here: a cost
    ratio measured under conditions this release does not define, and a count
    of avoided analyses that no committed run produces.
    """
    from docx import Document

    text = " ".join(p.text for p in Document(str(MANUSCRIPT)).paragraphs)
    for withdrawn in ("8x faster", "8× faster", "210 ", "400 avoided",
                      "1.4 nominal", "49 versus", "150-200 directions"):
        assert withdrawn not in text, f"withdrawn claim present: {withdrawn}"


def test_regenerates_identically_except_for_its_stamps(tmp_path):
    from docx import Document

    out = tmp_path / "rebuilt.docx"
    done = subprocess.run([sys.executable, str(BUILDER), "--out", str(out)],
                          capture_output=True, text=True, cwd=str(REPO_ROOT))
    assert done.returncode == 0, done.stdout + done.stderr
    original = [p.text for p in Document(str(MANUSCRIPT)).paragraphs]
    rebuilt = [p.text for p in Document(str(out)).paragraphs]
    # The closing note carries the commit, which moves between builds.
    assert original[:-1] == rebuilt[:-1]
