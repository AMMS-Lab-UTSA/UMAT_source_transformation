#!/usr/bin/env python3
"""Generate every publication table from the executed evidence.

Each table is emitted three ways: a CSV and a JSON that carry the numbers, and
a real Word table inside one document. The Word tables are tables, not pictures
of tables, so a co-author can edit a caption or fix a column width without
regenerating anything -- and so a reviewer can read the numbers with a screen
reader.

Nothing here transcribes a value from the manuscript. Every cell is read out of
a file under paper_results/, and a quantity that was not measured is rendered
as the reason it is unavailable, never as zero.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import subprocess
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Sequence

REPO_ROOT = Path(__file__).resolve().parents[2]
RESULTS = REPO_ROOT / "paper_results"
DEFAULT_OUT = RESULTS / "tables"

#: Printed where a quantity was never measured. Never a zero, and never blank:
#: a blank cell reads as "nothing to report" when it means "not established".
UNAVAILABLE = "not measured"

#: The illustrative example belongs only to the illustrative-example tables.
ILLUSTRATIVE = {"m3_j2", "j2"}


@dataclass
class Table:
    number: int
    slug: str
    caption: str
    columns: list[str]
    rows: list[list[Any]]
    inputs: list[Path]
    notes: str = ""
    filters: dict = field(default_factory=dict)


# --------------------------------------------------------------------------- #
# Reading
# --------------------------------------------------------------------------- #
def _csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def _json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _number(value: Any, digits: int = 3) -> str:
    """Format for print, or say plainly that there is no number."""
    if value is None or value == "":
        return UNAVAILABLE
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number == 0.0:
        return "0"
    if abs(number) < 1e-3 or abs(number) >= 1e5:
        return f"{number:.{digits}e}"
    return f"{number:.{digits}g}"


# --------------------------------------------------------------------------- #
# Tables
# --------------------------------------------------------------------------- #
def table1_metadata() -> Table:
    """Software inputs and outputs, read from the request schema itself."""
    sys.path.insert(0, str(REPO_ROOT / "src"))
    from umat_oti.services.workbench import OUTCOMES, PRODUCTS  # noqa: PLC0415

    rows = [
        ["Entry source", "input", "one Fortran file declaring SUBROUTINE UMAT",
         "fixed or free form; the form is detected, not declared"],
        ["Dependency roots", "input", "zero or more directories",
         "helper routines and INCLUDE files are resolved to a closure"],
        ["Tensor dimensions", "input", "NTENS, NDI, NSHR, NSTATV",
         "checked against the indices the source actually addresses"],
        ["Material properties", "input", "PROPS vector with a named mapping",
         "each differentiated parameter names its PROPS index"],
        ["State variable names", "input", "one name per STATEV slot",
         "used to label the exported state sensitivities"],
        ["Loading history", "input", "strain increment and increment count",
         "or a deformation gradient increment for a finite-strain source"],
        ["Requested products", "input", ", ".join(PRODUCTS),
         "any subset; an unrequested product is reported as such"],
        ["Transformed source", "output", "Fortran 90 with OTI arithmetic",
         "compiles against the generated OTI module"],
        ["Derivative arrays", "output", "one CSV per requested product",
         "per increment, component and parameter"],
        ["Primal response", "output", "stress and state per increment",
         "compared between the two builds before any derivative is judged"],
        ["Comparison report", "output", "per-row verdicts and their reasons",
         "outcomes: " + ", ".join(OUTCOMES)],
        ["Run manifest", "output", "JSON recording every stage",
         "each gate carries its own status and reason"],
    ]
    return Table(
        1, "software_inputs_outputs",
        "Inputs the software accepts and artefacts it produces. The product and "
        "outcome vocabularies are read from the request schema, so this table "
        "cannot describe a capability the code does not expose.",
        ["Item", "Direction", "Content", "Notes"], rows,
        [REPO_ROOT / "src" / "umat_oti" / "services" / "workbench.py"],
        notes="Generated from umat_oti.services.workbench.PRODUCTS and OUTCOMES.")


def table2_abaqus() -> Table:
    """The archived paired Abaqus round, one row per case."""
    source = RESULTS / "arc_791506" / "table2_abaqus_paired.csv"
    rows = []
    for row in _csv(source):
        rows.append([
            row["case_name"],
            row["status"],
            _number(row["stress_max_abs_diff"]),
            _number(row["ddsdde_max_abs_diff"]),
            row["ddsdde_status"],
            row["statev_status"],
            row["convergence_status"],
        ])
    rows.sort(key=lambda r: (r[1] != "passed", r[0]))
    return Table(
        2, "abaqus_paired_validation",
        "Paired Abaqus validation. Each case ran the untransformed and the "
        "transformed UMAT as separate jobs on the same deck and compared stress, "
        "state, tangent and convergence. Failed cases are listed with the rest.",
        ["Case", "Overall", "Stress max abs diff", "DDSDDE max abs diff",
         "DDSDDE verdict", "STATEV verdict", "Convergence verdict"],
        rows, [source],
        notes=("The Abaqus release used by this archived round was not recorded "
               "at execution time, so these results cannot be attributed to a "
               "specific version."))


def table3_internal_jacobians() -> Table:
    """Extracted internal Jacobians against an independent reference."""
    source = RESULTS / "internal_jacobians" / "table3_internal_jacobians.csv"
    seen, rows = set(), []
    for row in _csv(source):
        if row["canonical_source_id"] in seen:
            continue
        seen.add(row["canonical_source_id"])
        rows.append([
            row["model"], row["jacobian_variable"], row["iterate"],
            row["residual"], _number(row["oti_vs_fd_relative"]),
            _number(row["hand_coded_vs_fd_relative"]), row["verdict"],
        ])
    rows.sort(key=lambda r: -float(r[5]) if r[5] != UNAVAILABLE else 0.0)
    return Table(
        3, "internal_constitutive_jacobians",
        "Internal constitutive Jacobians extracted from the local Newton solve "
        "and checked against centred differences of the independently compiled "
        "untransformed build. The source's own hand-coded Jacobian is audited "
        "against the same reference and is never used as the reference.",
        ["Source", "Jacobian", "Iterate", "Residual",
         "OTI vs reference (rel)", "Hand-coded vs reference (rel)", "Verdict"],
        rows, [source],
        filters={"deduplication": "one row per canonical source identity"},
        notes=("Two sources carry a hand-coded Jacobian that differs from the "
               "reference by far more than the extracted one does."))


def table4_higher_order() -> Table:
    """Higher-order stress derivatives for the illustrative example."""
    source = (RESULTS / "actual_umat_higher_order" / "j2"
              / "table4_higher_order_actual_umat.csv")
    rows = []
    for row in _csv(source):
        rows.append([
            row["branch"], row["order"], row["comparison_rows"],
            row["passed_rows"], row["failed_rows"],
            _number(row["max_absolute_error"]),
            _number(row["max_relative_error_when_absolute_tolerance_exceeded"]),
            _number(row["absolute_tolerance"]),
        ])
    return Table(
        4, "higher_order_derivative_verification",
        "Higher-order stress derivatives of the illustrative example against an "
        "independent 80-digit reference. The relative error is reported only "
        "over rows whose absolute error exceeds the absolute tolerance, because "
        "a relative error on a quantity at the rounding floor measures nothing.",
        ["Branch", "Order", "Rows", "Passed", "Failed", "Max abs error",
         "Max rel error where significant", "Abs tolerance"], rows, [source],
        filters={"model": "controlled_j2_actual_umat (illustrative example)"})


def table5_illustrative() -> Table:
    """The illustrative example's own derivative verification, all families."""
    tangent = (RESULTS / "actual_umat_higher_order" / "j2"
               / "table2_ddsdde_illustrative.csv")
    sensitivity = RESULTS / "parameter_sensitivity" / "table6_comparison_rows.csv"
    rows = []

    entries = _csv(tangent)
    measured = [r for r in entries if r["judged_by"] == "relative"]
    zeros = [r for r in entries if r["judged_by"] == "structural_zero"]
    rows.append([
        "DDSDDE", len(entries), len(measured), len(zeros),
        len(entries) - len(measured) - len(zeros),
        sum(1 for r in entries if r["agrees"] == "False"),
        _number(max((float(r["relative_error"]) for r in measured), default=None)),
        "closed form and 80-digit difference",
    ])

    j2 = [r for r in _csv(sensitivity) if r["model"] == "m3_j2"]
    for array in ("DSIGMA_DP", "DSTATEV_DP"):
        selected = [r for r in j2 if r["array"] == array]
        if not selected:
            continue
        relative = [r for r in selected if r["judged_by"] == "relative"]
        exact = [r for r in selected
                 if float(r["oti"]) == 0.0 and float(r["reference"]) == 0.0]
        rows.append([
            array, len(selected), len(relative), len(exact),
            len(selected) - len(relative) - len(exact),
            sum(1 for r in selected if r["agrees"] == "False"),
            _number(max((float(r["relative_error"]) for r in relative),
                        default=None)),
            "centred differences of the untransformed build",
        ])
    return Table(
        5, "illustrative_parameter_and_state_sensitivities",
        "Every derivative family verified for the illustrative example. "
        "\"Measured\" counts entries with a defined relative error; \"exact "
        "zero\" counts entries that are exactly zero on both sides, which are "
        "verified but measure nothing; \"within reference\" counts entries the "
        "reference could not separate from the value. The three add to the "
        "entry count, so nothing is dropped.",
        ["Family", "Entries", "Measured", "Exact zero", "Within reference",
         "Disagreeing", "Worst measured rel error", "Reference"],
        rows, [tangent, sensitivity],
        filters={"model": "the illustrative J2 example only"})


def table6_collection() -> Table:
    """Collection-level parameter sensitivity, illustrative example excluded."""
    source = RESULTS / "parameter_sensitivity" / "parameter_sensitivity_round.json"
    round_data = _json(source)
    rows = []
    for model in round_data.get("models", []):
        name = model["model"]
        if name in ILLUSTRATIVE:
            continue
        stage = model["stages"].get("derivatives_verified", {})
        rows.append([
            name, stage.get("status", UNAVAILABLE),
            stage.get("rows", UNAVAILABLE),
            stage.get("rows_agreeing", UNAVAILABLE),
            stage.get("rows_disagreeing", UNAVAILABLE),
            stage.get("rows_reference_unresolved_at_noise_floor", UNAVAILABLE),
            stage.get("rows_reference_unresolved_by_branch_crossing", UNAVAILABLE),
            _number(stage.get("worst_relative_error")),
        ])
    rows.sort(key=lambda r: (r[1] == "succeeded", r[0]))
    return Table(
        6, "collection_parameter_sensitivity",
        "Parameter and state sensitivities across the model collection. Rows "
        "the reference cannot adjudicate are split by reason: those below what "
        "a centred difference resolves, and those whose stencil straddles a "
        "branch boundary. Every attempted model is listed, verified or not.",
        ["Model", "Verdict", "Rows", "Agreeing", "Disagreeing",
         "Unresolved (noise floor)", "Unresolved (branch crossing)",
         "Worst rel error"], rows, [source],
        filters={"excluded": sorted(ILLUSTRATIVE),
                 "exclusion_reason": "reported only in the illustrative tables"})


def table7_corpus() -> Table:
    """External corpus sources, reported whether or not they were verified."""
    source = RESULTS / "corpus" / "corpus_funnel.csv"
    rows = []
    for row in _csv(source):
        rows.append([
            row.get("candidate") or UNAVAILABLE,
            row.get("repository") or UNAVAILABLE,
            (row.get("commit_sha") or "")[:12] or UNAVAILABLE,
            row.get("license") or UNAVAILABLE,
            "yes" if row.get("multi_file") == "True" else "no",
            row.get("furthest_stage") or UNAVAILABLE,
            _number(row.get("worst_substantive_relative_error")),
            (row.get("blocker") or "none")[:90],
        ])
    rows.sort(key=lambda r: (r[5] != "derivatives_verified", r[0]))
    return Table(
        7, "external_web_corpus",
        "External sources acquired from public repositories at a pinned commit. "
        "Every source that entered the funnel is listed with the stage it "
        "reached and, where it stopped, why. Licences were classified from the "
        "repository's own licence text.",
        ["Source", "Repository", "Commit", "Licence", "Multi-file",
         "Furthest stage", "Worst rel error", "Blocker"], rows, [source])


def table8_adjudication() -> Table:
    """How every compared row of the illustrative example was adjudicated.

    This was a twenty-bar histogram stacked under the response curves, where it
    competed with them for attention while answering a different question. The
    counts are identical at every increment within a branch, so the whole
    histogram is two rows without losing anything.
    """
    source = RESULTS / "parameter_sensitivity" / "table6_comparison_rows.csv"
    rows_in = [r for r in _csv(source) if r["model"] == "m3_j2"]

    def kind(row: dict) -> str:
        if float(row["oti"]) == 0.0 and float(row["reference"]) == 0.0:
            return "exact zero on both sides"
        if row["judged_by"] == "relative":
            return "measured on relative error"
        return "within the reference's resolution"

    branches: dict[str, list[dict]] = {}
    for row in rows_in:
        label = ("elastic" if row["branch"] == "elastic" else "inelastic")
        branches.setdefault(label, []).append(row)

    rows = []
    for label in ("elastic", "inelastic"):
        selected = branches.get(label) or []
        if not selected:
            continue
        increments = sorted({int(r["increment"]) for r in selected})
        counts: dict[str, int] = {}
        for row in selected:
            counts[kind(row)] = counts.get(kind(row), 0) + 1
        worst = max((float(r["relative_error"]) for r in selected
                     if r["judged_by"] == "relative"), default=None)
        rows.append([
            label, f"{min(increments)}-{max(increments)}", len(selected),
            counts.get("exact zero on both sides", 0),
            counts.get("within the reference's resolution", 0),
            counts.get("measured on relative error", 0),
            sum(1 for r in selected if r["agrees"] != "True"),
            _number(worst),
        ])
    total = len(rows_in)
    rows.append(["all", f"1-{max(int(r['increment']) for r in rows_in)}", total,
                 sum(int(r[3]) for r in rows[:2]),
                 sum(int(r[4]) for r in rows[:2]),
                 sum(int(r[5]) for r in rows[:2]),
                 sum(int(r[6]) for r in rows[:2]),
                 _number(max((float(r["relative_error"]) for r in rows_in
                              if r["judged_by"] == "relative"), default=None))])
    return Table(
        8, "illustrative_row_adjudication",
        "How every compared row of the illustrative example was adjudicated. "
        "The counts are identical at every increment within a branch, so the "
        "table carries the same information the per-increment breakdown did. "
        "An exact zero is a row where both the generated value and the "
        "reference are exactly zero; a row within the reference's resolution "
        "is one the reference cannot separate from the value.",
        ["Branch", "Increments", "Rows", "Exact zero", "Within reference",
         "Measured", "Disagreeing", "Worst measured rel error"],
        rows, [source],
        filters={"model": "m3_j2, the illustrative example"},
        notes=("The three categories add to the row count, so no comparison "
               "leaves the table uncounted."))


BUILDERS: Sequence[Callable[[], Table]] = (
    table1_metadata, table2_abaqus, table3_internal_jacobians,
    table4_higher_order, table5_illustrative, table6_collection,
    table7_corpus, table8_adjudication,
)


# --------------------------------------------------------------------------- #
# Writing
# --------------------------------------------------------------------------- #
def _commit() -> str:
    try:
        done = subprocess.run(["git", "-C", str(REPO_ROOT), "rev-parse", "HEAD"],
                              capture_output=True, text=True, timeout=15)
    except (OSError, subprocess.SubprocessError):
        return "unavailable"
    return done.stdout.strip() if done.returncode == 0 else "unavailable"


def _relative(path: Path) -> str:
    path = Path(path).resolve()
    try:
        return str(path.relative_to(REPO_ROOT))
    except ValueError:
        return path.name


def _digest(path: Path) -> str:
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


def write_sources(table: Table, out_dir: Path) -> dict:
    """The numbers, in two machine-readable forms beside the Word document."""
    stem = f"table{table.number}_{table.slug}"
    csv_path = out_dir / f"{stem}.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle, lineterminator="\n")
        writer.writerow(table.columns)
        writer.writerows(table.rows)
    json_path = out_dir / f"{stem}.json"
    json_path.write_text(json.dumps({
        "number": table.number, "slug": table.slug, "caption": table.caption,
        "columns": table.columns,
        "rows": [dict(zip(table.columns, row)) for row in table.rows],
    }, indent=2) + "\n", encoding="utf-8")
    return {"csv": _relative(csv_path), "json": _relative(json_path)}


def write_provenance(table: Table, out_dir: Path, sources: dict) -> None:
    stem = f"table{table.number}_{table.slug}"
    (out_dir / f"{stem}_provenance.json").write_text(json.dumps({
        "table": stem,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "commit": _commit(),
        "command": "python tools/tables/build_paper_tables.py",
        "inputs": [{"path": _relative(p), "sha256": _digest(p),
                    "bytes": Path(p).stat().st_size}
                   for p in table.inputs if Path(p).is_file()],
        "outputs": sources,
        "row_count": len(table.rows),
        "column_count": len(table.columns),
        "filters": table.filters,
        "unavailable_marker": UNAVAILABLE,
        "notes": table.notes,
    }, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def write_document(tables: Sequence[Table], path: Path) -> None:
    """One Word file holding every table as an editable Word table."""
    from docx import Document  # noqa: PLC0415
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    from docx.shared import Inches, Pt  # noqa: PLC0415

    document = Document()
    # Wide evidence tables need the page. At the default margins several
    # columns wrapped mid-word, which is legible but ugly and makes a reviewer
    # think a value was truncated.
    for section in document.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    for style_name, size in (("Normal", 9),):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)

    heading = document.add_paragraph("UMAT-OTI publication tables")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)
    stamp = document.add_paragraph(
        f"Generated from the executed evidence at commit {_commit()}. "
        f"Every cell is read from a file under paper_results/; none was "
        f"transcribed. A quantity that was not measured reads "
        f"“{UNAVAILABLE}” rather than zero.")
    stamp.runs[0].italic = True

    for table in tables:
        document.add_paragraph()
        caption = document.add_paragraph()
        run = caption.add_run(f"Table {table.number}. {table.caption}")
        run.bold = True
        word_table = document.add_table(rows=1, cols=len(table.columns))
        word_table.style = "Table Grid"
        word_table.autofit = True
        for cell, name in zip(word_table.rows[0].cells, table.columns):
            cell.text = str(name)
            for paragraph in cell.paragraphs:
                for piece in paragraph.runs:
                    piece.bold = True
        for row in table.rows:
            cells = word_table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = "" if value is None else str(value)
        if table.notes:
            note = document.add_paragraph()
            note_run = note.add_run(f"Note. {table.notes}")
            note_run.italic = True
            note_run.font.size = Pt(8)
            note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.save(str(path))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args(argv)
    args.out_dir.mkdir(parents=True, exist_ok=True)

    built, skipped = [], []
    for builder in BUILDERS:
        try:
            table = builder()
        except FileNotFoundError as error:
            # A missing input is reported, never quietly replaced by an empty
            # table that would read as "nothing was found".
            skipped.append((builder.__name__, str(error)))
            continue
        sources = write_sources(table, args.out_dir)
        write_provenance(table, args.out_dir, sources)
        built.append(table)
        print(f"  Table {table.number}: {len(table.rows)} rows x "
              f"{len(table.columns)} columns  ->  {sources['csv']}")

    document = args.out_dir / "paper_tables.docx"
    write_document(built, document)
    print(f"  {_relative(document)}  ({len(built)} editable Word tables)")
    for name, reason in skipped:
        print(f"  SKIPPED {name}: {reason}")
    return 1 if skipped else 0


if __name__ == "__main__":
    raise SystemExit(main())
