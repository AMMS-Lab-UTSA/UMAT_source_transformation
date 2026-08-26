#!/usr/bin/env python3
"""Build the V5 SoftwareX manuscript from the executed evidence.

V4 is not read, edited or copied. This writes a new document whose every
numerical claim is substituted from `evidence_values.collect()`, so a claim the
evidence does not support cannot appear in the text: a missing key raises
rather than leaving a placeholder behind.

The figures embedded here are the real ones -- two screenshots of the working
interface and three plots rendered from committed evidence by committed
scripts. No conceptual diagram is included.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(Path(__file__).resolve().parent))
from evidence_values import collect  # noqa: E402

FIGURES = REPO_ROOT / "paper_results" / "figures"
DEFAULT_OUT = REPO_ROOT / "docs" / "manuscript" / "UMAT_OTI_SoftwareX_V5.docx"

#: SoftwareX limits the main text; the count excludes captions and references.
WORD_LIMIT = 3000

#: Where the parts of the interface a figure could not fit are reported.
OMITTED_ELSEWHERE = {
    "figure1_gui_request": "the run manifest",
    "figure2_gui_results": "Table 5 and the run manifest",
}


def _omission_sentence(stem: str) -> str:
    """Name what a GUI figure could not fit, from the capture's own record.

    Written from the provenance rather than by hand, so a caption cannot go on
    claiming to show a panel that a later capture had to drop.
    """
    record = FIGURES / "gui_screenshots_provenance.json"
    if not record.is_file():
        return ""
    regions = json.loads(record.read_text(encoding="utf-8")).get("regions", {})
    omitted = (regions.get(stem) or {}).get("omitted") or []
    if not omitted:
        return ""
    names = sorted(omitted)
    listed = (names[0] if len(names) == 1
              else " and ".join([", ".join(names[:-1]), names[-1]]))
    where = OMITTED_ELSEWHERE.get(stem, "the run manifest")
    plural = "panels are" if len(omitted) > 1 else "panel is"
    return (f" The {listed} {plural} not shown: the figure is cropped so its "
            f"text prints at readable size, and that content is reported in "
            f"{where}.")


FIGURE_FILES = [
    ("figure1_gui_request.png",
     "Figure 1. The interface constructing a request: the entry source and "
     "dependency roots, the tensor and state dimensions, the mapping of "
     "material parameters onto PROPS, the requested derivative products, and "
     "the supplied loading history."),
    ("figure2_gui_results.png",
     "Figure 2. The same interface after a real run. Every pipeline stage "
     "carries its own status; primal parity is reported above the derivative "
     "section; each product carries an outcome word rather than a colour. "
     "INTERNAL_JACOBIAN is reported unsupported for this model, with the "
     "reason, because it integrates its law without a local Newton iteration."),
    ("figure3_illustrative_derivatives.png",
     "Figure 3. Derivative validation for the illustrative example. (a) the "
     "consistent tangent against the closed-form reference; (b) repeated and "
     "(c) mixed higher-order stress derivatives against an independent "
     "80-digit reference; (d) where the error sits in each family relative to "
     "the agreement demanded of it."),
    ("figure4_parameter_sensitivities.png",
     "Figure 4. Parameter and state sensitivities over the complete loading "
     "path, each curve scaled by its own parameter so all stress curves carry "
     "MPa. (a) DSIGMA_DP; (b) DSTATEV_DP; (c) how every compared row was "
     "adjudicated, increment by increment."),
    ("figure5_collection_verification.png",
     "Figure 5. Verification across the deduplicated collection, excluding the "
     "illustrative example. (a) the offline route against the whole "
     "collection; (b) how many sources each route verified; (c) extracted and "
     "hand-coded internal Jacobians against the same independent reference."),
]

SECTIONS: list[tuple[str, list[str]]] = [
    ("Abstract", [
        "UMAT-OTI is an open-source source-to-source transformation tool that "
        "generates machine-precision derivatives of Abaqus user-material "
        "subroutines (UMATs) using order-truncated imaginary (OTI) arithmetic. "
        "Every derivative is specified by four elements -- a seed variable, a "
        "response variable, an extraction target and a differentiation order -- "
        "so one compact contract configures the consistent material tangent "
        "DDSDDE, the internal Jacobians of local constitutive solves, "
        "higher-order stress derivatives, and the exported point-wise "
        "parameter sensitivities DSIGMA_DP and DSTATEV_DP from a single "
        "unmodified UMAT. The transformed routine preserves the standard UMAT "
        "interface. Every product is checked against a reference that shares no "
        "code path with it: on an illustrative J2 model the generated tangent "
        "agrees with a closed-form consistent tangent over all "
        "{tangent_entries} entries to {tangent_worst_relative}, and derivatives "
        "through order {higher_order_max_order} agree with an independent "
        "80-digit reference over {higher_order_rows} comparisons with "
        "{higher_order_failed} failures. Across {unique_sources} "
        "globally deduplicated sources, {verified_either} are numerically "
        "verified by at least one route and {verified_neither} are not, each "
        "with a recorded reason. Extracting a model's internal Jacobian also "
        "audits the one it ships: in {jacobian_drifted} of {jacobian_sources} "
        "sources the hand-coded Jacobian differs from the independent "
        "reference by up to {jacobian_worst_hand_coded_percent}, while the "
        "extracted one agrees to {jacobian_worst_oti}."]),
    ("1. Motivation and significance", [
        "Abaqus user-material subroutines let researchers implement "
        "constitutive models absent from the standard libraries. A UMAT returns "
        "the updated stress, the updated internal variables and the material "
        "Jacobian DDSDDE, which in Abaqus/Standard is the consistent tangent of "
        "the global Newton-Raphson procedure. An inconsistent tangent raises "
        "cost, triggers cutbacks or fails outright.",
        "The need for accurate derivatives extends beyond DDSDDE. Many models "
        "contain nonlinear solution procedures inside the stress update itself, "
        "such as return mapping for plasticity and viscoplasticity, where a "
        "local residual is solved for an internal variable. These require "
        "internal Jacobians distinct from the tangent handed to Abaqus, so a "
        "practical differentiation tool must produce both.",
        "Derivatives are traditionally obtained by manual derivation or finite "
        "differences. Manual derivation is exact and remains common practice, "
        "but it is labour-intensive, difficult to maintain and prone to human "
        "error. Finite-difference approaches are easier to implement but "
        "require a careful choice of perturbation size, suffer truncation and "
        "round-off error, and grow costly with the number of variables and "
        "derivative orders.",
        "Automatic differentiation instead propagates derivative information "
        "through the computational graph of the implemented algorithm. "
        "Source-transformation tools such as Tapenade and ADIFOR, "
        "operator-overloading libraries such as Adept and Sacado, and "
        "compiler-based frameworks such as Enzyme have demonstrated these "
        "strategies, and the same ideas reached finite-element and "
        "constitutive-model development through AceGen, dolfin-adjoint, FEniCS "
        "and UMAT4COMSOL. Hypercomplex automatic differentiation encodes "
        "derivatives in the coefficients of extended-number algebras; "
        "complex-step, hyper-dual, multicomplex and OTI arithmetic all belong "
        "to this family, and none suffers the subtraction error of finite "
        "differences.",
        "Constitutive models are usually mature Fortran code bases that must "
        "keep the interface Abaqus expects, and their authors need derivatives "
        "at more than one level. UMAT-OTI applies OTI differentiation directly "
        "to that source, leaving the interface intact. The contribution is not "
        "a new differentiation method but a framework that adapts one to Abaqus "
        "UMAT development, together with the support library, reports and "
        "validation artefacts that make the generated code auditable."]),
    ("2. Software description", [
        "A request names what to perturb, what to watch, where to write the "
        "result and how many orders to keep. The same request object is "
        "constructed by the command line, by the batch runner and by the "
        "graphical interface, and all three call one execution path, so a "
        "result shown in the interface is produced by the code that produces "
        "the published evidence.",
        "The transformation resolves the entry routine's dependency closure -- "
        "helper subprograms, INCLUDE files and the modules they need -- lifts "
        "the arithmetic to the OTI type, generates the algebra module for the "
        "requested number of directions and order, and emits Fortran that "
        "compiles against it. {multi_file_sources} of the verified sources are "
        "multi-file: their closure is resolved rather than declared, and a "
        "helper defined differently in two places is reported as ambiguous "
        "instead of being chosen arbitrarily.",
        "Verification is gated. Source identity and closure come first, then "
        "transformation, then independent compilation of the original and the "
        "transformed routine, then execution of both, then primal parity, then "
        "resolution of the reference, and only then the derivative comparison. "
        "A case failing an earlier gate contributes no verified derivative "
        "rows, and compiling is never reported as verification. Figures 1 and 2 "
        "show the interface constructing a request and reporting one such run."]),
    ("3. Illustrative example", [
        "One three-dimensional small-strain J2 plasticity model with linear "
        "isotropic hardening runs through every derivative product. It has six "
        "stress components and one state variable, the accumulated equivalent "
        "plastic strain, which makes it history dependent: the stress at the "
        "end of an increment depends on the plastic strain inherited from every "
        "increment before it. The loading path used throughout begins elastic "
        "and yields at increment {j2_yield_increment} of {j2_increments}.",
        "The consistent tangent is checked against two references that fail "
        "differently: the closed-form elastoplastic consistent tangent, and an "
        "80-digit centred difference of an independent integrator. Over "
        "{tangent_entries} entries none disagrees. {tangent_structural_zeros} "
        "are zeros of the matrix, returned as exactly zero; the remaining "
        "{tangent_measured} agree to {tangent_worst_relative}, while the two "
        "references agree with each other to {tangent_reference_spread}. "
        "Raising the requested order on the same path retains the higher "
        "derivatives of the same stress update: {higher_order_rows} comparisons "
        "through order {higher_order_max_order}, {higher_order_failed} "
        "failures, worst relative difference "
        "{higher_order_worst_relative} where the quantity is large enough for a "
        "relative error to mean anything. Figure 3 collects these.",
        "The parameter sensitivities perturb the material constants rather than "
        "the strain. DSIGMA_DP holds the derivative of each stress component "
        "with respect to each material parameter and DSTATEV_DP the derivative "
        "of the state variable, both at one material point and one increment. "
        "The second is needed because the plastic strain carried into an "
        "increment already depends on the parameters. Over the full path, "
        "{j2_sensitivity_rows} comparisons agree with centred differences of "
        "the independently compiled original, {j2_sensitivity_disagreeing} "
        "disagree, and the worst relative difference is "
        "{j2_sensitivity_worst_relative}. {j2_exact_zeros} of those rows are "
        "exactly zero on both sides: before yield the stress does not depend on "
        "the yield stress or the hardening modulus, and the equivalent plastic "
        "strain depends on nothing. Those are structural zeros of the model, "
        "read from the raw values rather than inferred from a plot. Figure 4 "
        "shows both arrays over the path."]),
    ("4. Verification across the collection", [
        "Sources are counted after global identity reconciliation. A UMAT "
        "reachable both from the in-repository archive and from a pinned "
        "upstream snapshot is one implementation with two origins, not two "
        "sources: {raw_discovered_files} discovered files reduce to "
        "{deduplicated_sources} distinct implementations across "
        "{upstream_repositories} independent upstream repositories, spanning "
        "{constitutive_models} constitutive models. The illustrative example is "
        "excluded from every count in this section.",
        "Of {unique_sources} collection sources, {verified_offline} are "
        "verified offline -- transformed, compiled twice, checked for primal "
        "parity and compared against an independently replayed reference -- and "
        "{verified_abaqus} passed a paired Abaqus round in which the original "
        "and transformed subroutines ran as separate jobs on the same deck. "
        "{verified_either} are verified by at least one route. The remaining "
        "{verified_neither} are kept in the denominator with a recorded reason: "
        "an ambiguous helper closure, an absent upstream property vector, a "
        "failed execution in the archived Abaqus round, and two whose reference "
        "cannot adjudicate every row.",
        "Across {collection_models} models the sweep produced "
        "{collection_rows} comparisons, {collection_agreeing} agreeing and "
        "{collection_disagreeing} disagreeing, with a worst relative difference "
        "of {collection_worst_relative}. The {collection_unresolved} rows the "
        "reference cannot settle are reported by reason rather than pooled: "
        "{collection_noise_floor} sit below what a centred difference resolves "
        "at any step, and {collection_branch_crossing} sit on the increment "
        "where a Drucker-Prager model first yields, where the stencil straddles "
        "the kink and returns a secant across it rather than the derivative on "
        "the branch the increment took. Those rows withhold their directions "
        "rather than being counted either way.",
        "Extracting a model's internal Jacobian also audits the one it ships. "
        "Across {jacobian_sources} sources carrying a local Newton solve, the "
        "extracted coefficient agrees with centred differences of the "
        "independently compiled original to {jacobian_worst_oti}. The "
        "hand-coded coefficient is measured against the same reference and is "
        "never used as it. In {jacobian_drifted} sources it differs by far "
        "more: {jacobian_worst_hand_coded_percent} in "
        "{jacobian_worst_hand_coded_model} and "
        "{jacobian_second_hand_coded_percent} in "
        "{jacobian_second_hand_coded_model}. Both models converge and produce "
        "plausible stress, so the drift is invisible without a derivative that "
        "follows the implementation. Figure 5 collects the collection results."]),
    ("5. Impact", [
        "UMAT-OTI changes how derivative implementations are maintained. A "
        "change to a material model traditionally requires a fresh manual "
        "derivation of the tangent and fresh verification of the resulting "
        "code. Here the implemented stress update remains the primary "
        "representation of the model and the derivatives follow from it "
        "automatically, which extends the useful lifetime of existing UMATs.",
        "The products serve different consumers. Accurate tangents improve the "
        "robustness of implicit finite-element analysis, internal Jacobians "
        "support local return mapping and viscoplastic integration, and "
        "higher-order derivatives supply curvature to calibration and to solver "
        "strategies that fit more than a straight line. The parameter "
        "sensitivities give calibration, uncertainty propagation, inverse "
        "identification and surrogate modelling a point-wise derivative of the "
        "constitutive response with a documented shape and parameter ordering.",
        "A second contribution is transparency. Every transformation produces "
        "the transformed source, the support library, a readable report and "
        "optional validation outputs, so the differentiation can be audited and "
        "reproduced. A difference against a hand-coded reference is not "
        "automatically a transformation error: several of them exposed "
        "approximate tangents, drifted internal Jacobians or outright errors in "
        "the original UMAT, which makes the software a verification aid as much "
        "as a differentiation tool."]),
    ("6. Conclusions", [
        "UMAT-OTI applies order-truncated imaginary arithmetic to existing "
        "Abaqus UMAT implementations, preserving the standard interface while "
        "generating machine-precision derivatives from the original stress "
        "update. One contract, naming a seed, a response, a target and an "
        "order, configures the material tangent, the internal constitutive "
        "Jacobians, the higher-order stress derivatives and the two exported "
        "sensitivities.",
        "Each product was verified against a reference sharing no code path "
        "with it, and every case that was attempted is reported whether or not "
        "it succeeded. The implementation targets UMAT code paths that can be "
        "analysed and transformed automatically; code relying on external "
        "helpers, on solver-provided storage, or on library routines that "
        "cannot operate on the extended type needs manual preparation first, "
        "and the sources where that applies are named in the collection "
        "evidence rather than omitted from it."]),
    ("Acknowledgements", [
        "The authors gratefully acknowledge financial support from the National "
        "Aeronautics and Space Administration under Grant No. 80NSSC23K1342 and "
        "from the National Science Foundation through CAREER Award "
        "No. 2237313."]),
]


class StrictValues(dict):
    """A substitution map that refuses to leave a claim unsupported.

    A KeyError here means the text asserts something the evidence does not
    provide. Falling back to the literal placeholder would put "{...}" into a
    manuscript; falling back to a blank would put an unsupported sentence into
    one, which is worse.
    """

    def __missing__(self, key: str):
        raise KeyError(
            f"the manuscript cites {key!r}, which evidence_values.collect() "
            "does not provide. Add the measurement or remove the claim.")


def _derived(values: dict) -> dict[str, str]:
    """Values the text states in a different form from the evidence."""
    text = {key: value.text() for key, value in values.items()}
    for key in ("jacobian_worst_hand_coded", "jacobian_second_hand_coded"):
        text[f"{key}_percent"] = f"{values[key].value * 100:.3g}%"
    text["collection_unresolved"] = str(
        values["collection_noise_floor"].value
        + values["collection_branch_crossing"].value)
    return text


def _fill(paragraph: str, substitutions: dict[str, str]) -> str:
    return paragraph.format_map(StrictValues(substitutions))


def _commit() -> str:
    try:
        done = subprocess.run(["git", "-C", str(REPO_ROOT), "rev-parse", "HEAD"],
                              capture_output=True, text=True, timeout=15)
    except (OSError, subprocess.SubprocessError):
        return "unavailable"
    return done.stdout.strip() if done.returncode == 0 else "unavailable"


def _word_count(paragraphs: list[str]) -> int:
    return sum(len(re.findall(r"\S+", text)) for text in paragraphs)


def build(out_path: Path) -> dict:
    from docx import Document  # noqa: PLC0415
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    from docx.shared import Inches, Pt  # noqa: PLC0415

    values = collect()
    substitutions = _derived(values)

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    run = title.add_run("UMAT-OTI: Automatic generation of material and "
                        "constitutive Jacobians for Abaqus UMATs")
    run.bold = True
    run.font.size = Pt(15)

    body: list[str] = []
    for heading, paragraphs in SECTIONS:
        head = document.add_paragraph()
        head_run = head.add_run(heading)
        head_run.bold = True
        head_run.font.size = Pt(12)
        for paragraph in paragraphs:
            filled = _fill(paragraph, substitutions)
            body.append(filled)
            _add_text(document.add_paragraph(), filled)

        if heading.startswith("2."):
            _add_figures(document, FIGURE_FILES[:2], Inches, Pt,
                         WD_ALIGN_PARAGRAPH)
        elif heading.startswith("3."):
            _add_figures(document, FIGURE_FILES[2:4], Inches, Pt,
                         WD_ALIGN_PARAGRAPH)
        elif heading.startswith("4."):
            _add_figures(document, FIGURE_FILES[4:], Inches, Pt,
                         WD_ALIGN_PARAGRAPH)

    words = _word_count(body)
    document.add_paragraph()
    note = document.add_paragraph()
    note_run = note.add_run(
        f"Every numerical value in this manuscript was substituted from the "
        f"executed evidence at commit {_commit()} by "
        f"tools/manuscript/build_v5_manuscript.py. Tables are generated "
        f"separately into paper_results/tables/paper_tables.docx. Main text: "
        f"{words} words of a {WORD_LIMIT}-word limit.")
    note_run.italic = True
    note_run.font.size = Pt(9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return {"words": words, "figures": len(FIGURE_FILES), "values": values,
            "body": body}


#: "1.23x10^-15" as the evidence formats it, so the exponent can become a real
#: superscript instead of printing a caret in a manuscript.
_EXPONENT = re.compile(r"(×10)\^(-?\d+)")


def _add_text(paragraph, text: str):
    """Write a paragraph, rendering exponents as superscripts."""
    text = text.replace(" -- ", " \u2014 ")
    position = 0
    for match in _EXPONENT.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        paragraph.add_run(match.group(1))
        exponent = paragraph.add_run(match.group(2))
        exponent.font.superscript = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])
    return paragraph


def _add_figures(document, entries, Inches, Pt, alignment) -> None:
    for name, caption in entries:
        path = FIGURES / name
        if not path.is_file():
            raise FileNotFoundError(
                f"{name} has not been generated; run the figure scripts before "
                "building the manuscript rather than shipping it without them")
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(6.2))
        caption = caption + _omission_sentence(path.stem)
        caption_paragraph = _add_text(document.add_paragraph(), caption)
        for caption_run in caption_paragraph.runs:
            caption_run.font.size = Pt(9)
            caption_run.italic = True


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args(argv)

    result = build(args.out)
    values = result["values"]
    out_path = args.out.resolve()
    try:
        shown = out_path.relative_to(REPO_ROOT)
    except ValueError:
        # An out-of-tree build, as a test does. Naming the absolute path in a
        # record that may be committed is what the sanitiser elsewhere exists
        # to prevent, so the file name alone is recorded.
        shown = Path(out_path.name)
    provenance = args.out.with_name(f"{args.out.stem}_provenance.json")
    provenance.write_text(json.dumps({
        "manuscript": str(shown),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "commit": _commit(),
        "command": "python tools/manuscript/build_v5_manuscript.py",
        "word_count": result["words"],
        "word_limit": WORD_LIMIT,
        "figure_count": result["figures"],
        "figures": [name for name, _ in FIGURE_FILES],
        "substituted_values": {
            key: {"value": value.value, "source": value.source}
            for key, value in sorted(values.items())},
        "sha256": hashlib.sha256(args.out.read_bytes()).hexdigest(),
        "note": ("V4 is not read, edited or copied by this script. Every "
                 "numerical claim is substituted from the evidence; a claim "
                 "the evidence does not support raises rather than producing "
                 "a placeholder."),
    }, indent=2, sort_keys=True, default=str) + "\n", encoding="utf-8")

    print(f"  {shown}")
    print(f"  {result['words']} words of {WORD_LIMIT}; "
          f"{result['figures']} figures")
    print(f"  {len(values)} values substituted from evidence")
    if result["words"] > WORD_LIMIT:
        print(f"  OVER the SoftwareX limit by {result['words'] - WORD_LIMIT} "
              "words")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
